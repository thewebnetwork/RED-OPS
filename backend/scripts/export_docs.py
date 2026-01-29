#!/usr/bin/env python3
"""
Export System Documentation to PDF and DOCX formats
"""

import markdown2
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import os

# Input/Output paths
INPUT_MD = "/app/backups/system_docs/Red_Pulse_System_Documentation.md"
OUTPUT_DOCX = "/app/backups/system_docs/Red_Pulse_System_Documentation.docx"
OUTPUT_PDF = "/app/backups/system_docs/Red_Pulse_System_Documentation.pdf"
OUTPUT_HTML = "/app/backups/system_docs/Red_Pulse_System_Documentation.html"

def read_markdown():
    """Read the markdown file"""
    with open(INPUT_MD, 'r', encoding='utf-8') as f:
        return f.read()

def create_docx(md_content):
    """Create DOCX from markdown content"""
    doc = Document()
    
    # Set document properties
    core_props = doc.core_properties
    core_props.author = "Red Pulse System"
    core_props.title = "Red Pulse System Documentation Pack"
    
    # Process markdown content
    lines = md_content.split('\n')
    
    current_table = []
    in_table = False
    in_code_block = False
    code_block_content = []
    
    for line in lines:
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_content)
                p = doc.add_paragraph()
                p.style = 'Normal'
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                code_block_content = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue
        
        # Handle tables
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                current_table = []
            
            # Skip separator rows
            if '---' in line:
                continue
            
            # Parse table row
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                current_table.append(cells)
            continue
        elif in_table:
            # End of table - create it
            if current_table:
                try:
                    # Create table
                    num_cols = len(current_table[0])
                    table = doc.add_table(rows=len(current_table), cols=num_cols)
                    table.style = 'Table Grid'
                    
                    for i, row_data in enumerate(current_table):
                        row = table.rows[i]
                        for j, cell_text in enumerate(row_data[:num_cols]):
                            cell = row.cells[j]
                            # Clean markdown formatting
                            clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                            clean_text = re.sub(r'\*([^*]+)\*', r'\1', clean_text)
                            clean_text = re.sub(r'`([^`]+)`', r'\1', clean_text)
                            cell.text = clean_text
                            
                            # Bold header row
                            if i == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                except Exception as e:
                    print(f"Table error: {e}")
                
                current_table = []
            in_table = False
        
        # Handle headings
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:].strip(), level=3)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            # Bullet point
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold/italic in list items
            add_formatted_text(p, text)
        elif line.strip().startswith('1. ') or re.match(r'^\d+\.', line.strip()):
            # Numbered list
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
        elif line.strip() == '---':
            # Horizontal rule - skip or add page break
            pass
        elif line.strip():
            # Regular paragraph
            p = doc.add_paragraph()
            add_formatted_text(p, line.strip())
    
    # Save document
    doc.save(OUTPUT_DOCX)
    print(f"✅ DOCX created: {OUTPUT_DOCX}")

def add_formatted_text(paragraph, text):
    """Add text with basic formatting to a paragraph"""
    # Simple approach: just add the text with basic cleanup
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Remove bold markers
    text = re.sub(r'\*([^*]+)\*', r'\1', text)      # Remove italic markers
    text = re.sub(r'`([^`]+)`', r'\1', text)        # Remove code markers
    paragraph.add_run(text)

def create_html(md_content):
    """Create HTML from markdown for PDF conversion"""
    html_content = markdown2.markdown(
        md_content, 
        extras=['tables', 'fenced-code-blocks', 'header-ids', 'code-friendly']
    )
    
    # Create full HTML document with styling
    full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Red Pulse System Documentation Pack</title>
    <style>
        body {{
            font-family: 'Segoe UI', Arial, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px;
            color: #333;
        }}
        h1 {{
            color: #A2182C;
            border-bottom: 2px solid #A2182C;
            padding-bottom: 10px;
            page-break-before: always;
        }}
        h1:first-of-type {{
            page-break-before: avoid;
        }}
        h2 {{
            color: #333;
            border-bottom: 1px solid #ddd;
            padding-bottom: 5px;
            margin-top: 30px;
        }}
        h3 {{
            color: #555;
            margin-top: 25px;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
            font-size: 14px;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 10px;
            text-align: left;
        }}
        th {{
            background-color: #f5f5f5;
            font-weight: bold;
        }}
        tr:nth-child(even) {{
            background-color: #fafafa;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Consolas', monospace;
            font-size: 13px;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
            font-size: 13px;
        }}
        pre code {{
            background: none;
            padding: 0;
        }}
        ul, ol {{
            padding-left: 25px;
        }}
        li {{
            margin-bottom: 5px;
        }}
        .toc {{
            background: #f9f9f9;
            padding: 20px;
            border-radius: 5px;
            margin-bottom: 30px;
        }}
        .toc a {{
            color: #A2182C;
            text-decoration: none;
        }}
        .toc a:hover {{
            text-decoration: underline;
        }}
        @media print {{
            body {{
                padding: 20px;
            }}
            h1 {{
                page-break-before: always;
            }}
            h1:first-of-type {{
                page-break-before: avoid;
            }}
            table {{
                page-break-inside: avoid;
            }}
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
    
    # Save HTML
    with open(OUTPUT_HTML, 'w', encoding='utf-8') as f:
        f.write(full_html)
    print(f"✅ HTML created: {OUTPUT_HTML}")
    
    return full_html

def create_pdf(html_content):
    """Create PDF from HTML using WeasyPrint"""
    try:
        from weasyprint import HTML
        HTML(string=html_content).write_pdf(OUTPUT_PDF)
        print(f"✅ PDF created: {OUTPUT_PDF}")
    except Exception as e:
        print(f"⚠️ PDF creation failed: {e}")
        print("   Try opening the HTML file in a browser and print to PDF")

def main():
    print("=" * 60)
    print("RED PULSE SYSTEM DOCUMENTATION EXPORT")
    print("=" * 60)
    
    # Read markdown
    print("\n📖 Reading markdown source...")
    md_content = read_markdown()
    print(f"   Read {len(md_content)} characters")
    
    # Create DOCX
    print("\n📄 Creating DOCX...")
    create_docx(md_content)
    
    # Create HTML
    print("\n🌐 Creating HTML...")
    html_content = create_html(md_content)
    
    # Create PDF
    print("\n📑 Creating PDF...")
    create_pdf(html_content)
    
    print("\n" + "=" * 60)
    print("EXPORT COMPLETE")
    print("=" * 60)
    print(f"\nFiles created in: /app/backups/system_docs/")
    print(f"  - Red_Pulse_System_Documentation.md   (Markdown)")
    print(f"  - Red_Pulse_System_Documentation.docx (Word)")
    print(f"  - Red_Pulse_System_Documentation.html (HTML)")
    print(f"  - Red_Pulse_System_Documentation.pdf  (PDF)")

if __name__ == "__main__":
    main()
